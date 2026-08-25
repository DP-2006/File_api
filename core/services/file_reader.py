# core/services/file_reader.py

import os
import zipfile
import tarfile
import tempfile
import subprocess
from django.core.files.uploadedfile import UploadedFile as DjangoUploadedFile

class FileReader:
    """خواننده حرفه‌ای فایل‌های مختلف"""
    
    @staticmethod
    def read_file(file_obj) -> dict:
        """خواندن فایل و استخراج محتوا بر اساس نوع"""
        file_path = file_obj.path if hasattr(file_obj, 'path') else None
        file_name = file_obj.name if hasattr(file_obj, 'name') else str(file_obj)
        ext = os.path.splitext(file_name)[1].lower()
        
        result = {
            'filename': file_name,
            'extension': ext,
            'content': '',
            'metadata': {},
            'error': None
        }
        
        if ext == '.txt':
            result.update(FileReader._read_text(file_path))
        elif ext == '.pdf':
            result.update(FileReader._read_pdf(file_path))
        elif ext == '.docx':
            result.update(FileReader._read_docx(file_path))
        elif ext in ['.xlsx', '.xls']:
            result.update(FileReader._read_excel(file_path))
        elif ext in ['.jar', '.java']:
            result.update(FileReader._read_jar(file_path))
        elif ext in ['.rar', '.zip', '.tar', '.gz']:
            result.update(FileReader._read_archive(file_path, ext))
        else:
            result['content'] = f"فایل {ext} - فقط نام فایل قابل خواندن است"
        
        return result


    @staticmethod
    def _read_text(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return {'content': content[:10000], 'metadata': {'size': len(content)}}
        except:
            return {'content': 'خطا در خواندن فایل متنی', 'metadata': {}}



    @staticmethod
    def _read_pdf(file_path):
        try:
            import PyPDF2
            content = ""
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                for page in pdf.pages[:20]:  # حداکثر 20 صفحه
                    content += page.extract_text() or ""
            return {'content': content[:15000], 'metadata': {'pages': len(pdf.pages)}}
        except Exception as e:
            return {'content': f'خطا در خواندن PDF: {e}', 'metadata': {}}
    

    @staticmethod
    def _read_docx(file_path):
        try:
            import docx
            doc = docx.Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            return {'content': content[:10000], 'metadata': {'paragraphs': len(doc.paragraphs)}}
        except Exception as e:
            return {'content': f'خطا در خواندن DOCX: {e}', 'metadata': {}}
      


    @staticmethod
    def _read_excel(file_path):
        try:
            import pandas as pd
            content = ""
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names[:3]:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                content += f"\n\nصفحه: {sheet_name}\n"
                content += df.to_string()[:3000]
            return {'content': content[:10000], 'metadata': {'sheets': excel_file.sheet_names}}
        except Exception as e:
            return {'content': f'خطا در خواندن Excel: {e}', 'metadata': {}}
    
    @staticmethod
    def _read_jar(file_path):
        """خواندن فایل JAR (مشابه ZIP)"""
        try:
            content = ""
            with zipfile.ZipFile(file_path, 'r') as jar:
                files_list = jar.namelist()[:20]
                content = f"فایل‌های داخل JAR: {', '.join(files_list)}\n\n"
                # خواندن فایل‌های متنی داخل JAR
                for fname in files_list:
                    if fname.endswith(('.java', '.xml', '.properties', '.txt')):
                        try:
                            file_content = jar.read(fname).decode('utf-8', errors='ignore')[:500]
                            content += f"\n--- {fname} ---\n{file_content}\n"
                        except:
                            pass
            return {'content': content[:10000], 'metadata': {'type': 'jar'}}
        except Exception as e:
            return {'content': f'خطا در خواندن JAR: {e}', 'metadata': {}}
    
    @staticmethod
    def _read_archive(file_path, ext):
        """خواندن فایل‌های فشرده (ZIP, RAR, TAR)"""
        content = ""
        temp_dir = tempfile.mkdtemp()
        
        try:
            if ext == '.zip':
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    files_list = zip_ref.namelist()[:30]
                    content = f"فایل‌های داخل ZIP: {', '.join(files_list)}\n\n"
                    
                    # تلاش برای خواندن فایل‌های متنی داخل ZIP
                    for fname in files_list:
                        if fname.endswith(('.txt', '.py', '.java', '.xml', '.json', '.md')):
                            try:
                                file_content = zip_ref.read(fname).decode('utf-8', errors='ignore')[:500]
                                content += f"\n--- {fname} ---\n{file_content}\n"
                            except:
                                pass
                            
            elif ext == '.rar':
                try:
                    import rarfile
                    with rarfile.RarFile(file_path) as rar:
                        files_list = rar.namelist()[:30]
                        content = f"فایل‌های داخل RAR: {', '.join(files_list)}\n\n"
                except:
                    content = "فایل RAR - برای خواندن نیاز به نصب rarfile دارد"
                    
            elif ext in ['.tar', '.gz']:
                with tarfile.open(file_path, 'r:*') as tar:
                    files_list = tar.getnames()[:30]
                    content = f"فایل‌های داخل TAR: {', '.join(files_list)}\n\n"
            
            return {'content': content[:10000], 'metadata': {'archive_type': ext}}
        except Exception as e:
            return {'content': f'خطا در خواندن آرشیو: {e}', 'metadata': {}}


# نصب کتابخانه‌ها در صورت نیاز
try:
    import rarfile
except ImportError:
    pass  # اختیاری